import os
import pdfplumber
import docx
import tempfile
from collections import defaultdict
import re


def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(x_tolerance=2, y_tolerance=2, keep_blank_chars=False)
            if not words:
                text += page.extract_text() or ""
                continue

            # Column detection
            words_sorted = sorted(words, key=lambda w: w['x0'])
            column_edges = []
            column_gap_threshold = 40
            prev_x0 = None
            for w in words_sorted:
                x0 = w['x0']
                if prev_x0 is None or abs(x0 - prev_x0) > column_gap_threshold:
                    column_edges.append(x0)
                prev_x0 = x0
            column_edges = sorted(set(column_edges))

            def find_column(x0):
                for i, edge in enumerate(column_edges):
                    if x0 < edge + column_gap_threshold:
                        return i
                return len(column_edges) - 1

            for w in words:
                w['column'] = find_column(w['x0'])

            # Sort and group by columns
            columns = {}
            for w in words:
                columns.setdefault(w['column'], []).append(w)

            column_indices = sorted(columns.keys(), key=lambda c: column_edges[c])
            page_text = ""

            for col_idx in column_indices:
                col_words = columns[col_idx]
                col_words_sorted = sorted(col_words, key=lambda w: (w['top'], w['x0']))
                lines = group_by_line(col_words_sorted, tolerance=3)
                lines_sorted = sorted(lines.items(), key=lambda x: x[0])

                for _, line_words in lines_sorted:
                    line = " ".join(w['text'] for w in sorted(line_words, key=lambda w: w['x0']))
                    page_text += line.strip() + "\n"

            text += page_text + "\n"
    return text.strip()


def group_by_line(words, tolerance=3):
    lines = defaultdict(list)
    for word in words:
        added = False
        for top in lines:
            if abs(word['top'] - top) <= tolerance:
                lines[top].append(word)
                added = True
                break
        if not added:
            lines[word['top']].append(word)
    return lines


def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)
    return "\n".join(full_text).strip()


def fix_broken_words(text):
    """
    Fix uppercase words that were spaced out like 'E D U C A T I O N'.
    Detects lines made up of single letters separated by spaces and rejoins them.
    """
    fixed_lines = []
    for line in text.splitlines():
        if re.fullmatch(r"([A-Z]\s+){2,}[A-Z]", line.strip()):
            fixed_lines.append(line.replace(" ", ""))
        else:
            fixed_lines.append(line)
    return "\n".join(fixed_lines)


def format_section_headers(text):
    """
    Adds newlines before likely section headers for better readability.
    """
    lines = text.splitlines()
    updated = []
    for line in lines:
        if re.fullmatch(r"[A-Z][A-Z\s\-]{2,}", line.strip()):
            updated.append("")  # newline before headers
        updated.append(line)
    return "\n".join(updated)


def enforce_line_structure(text):
    """
    Ensures that bullets and common list formats appear on new lines.
    """
    text = re.sub(r"(?<!\n)(- )", r"\n\1", text)  # bullet points
    text = re.sub(r"\n{3,}", "\n\n", text)  # no extra empty lines
    return text.strip()


def parse_resume(uploaded_file):
    file_ext = os.path.splitext(uploaded_file.name)[-1].lower()
    suffix = file_ext if file_ext in [".pdf", ".docx"] else ""

    fd, temp_path = tempfile.mkstemp(suffix=suffix)
    try:
        with os.fdopen(fd, "wb") as tmp:
            tmp.write(uploaded_file.read())

        if file_ext == ".pdf":
            raw_text = extract_text_from_pdf(temp_path)
        elif file_ext == ".docx":
            raw_text = extract_text_from_docx(temp_path)
        else:
            return "Unsupported file type. Please upload a PDF or DOCX."

        # Dynamic post-processing
        text = fix_broken_words(raw_text)
        text = format_section_headers(text)
        text = enforce_line_structure(text)

        return text

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
